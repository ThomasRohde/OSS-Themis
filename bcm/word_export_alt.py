from typing import List
from bcm.models import LayoutModel
from bcm.layout_manager import process_layout
from bcm.settings import Settings
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
import re
import datetime

def _add_inline_formatting(paragraph, text: str) -> None:
    """Add inline formatting (bold, italic) to a paragraph."""
    segments = []
    current_pos = 0
    text_length = len(text)
    bold_delimiters = ('**', '__')
    italic_delimiters = ('*', '_')

    while current_pos < text_length:
        found = False
        # Check for bold delimiters
        for delim in bold_delimiters:
            if text.startswith(delim, current_pos):
                end_pos = current_pos + len(delim)
                close_pos = text.find(delim, end_pos)
                if close_pos != -1:
                    if current_pos > 0:
                        segments.append(('normal', text[current_pos:end_pos - len(delim)]))
                    bold_text = text[end_pos:close_pos]
                    segments.append(('bold', bold_text))
                    current_pos = close_pos + len(delim)
                    found = True
                    break
        if found:
            continue

        # Check for italic delimiters
        for delim in italic_delimiters:
            if text.startswith(delim, current_pos):
                end_pos = current_pos + len(delim)
                close_pos = text.find(delim, end_pos)
                if close_pos != -1:
                    if current_pos > 0:
                        segments.append(('normal', text[current_pos:end_pos - len(delim)]))
                    italic_text = text[end_pos:close_pos]
                    segments.append(('italic', italic_text))
                    current_pos = close_pos + len(delim)
                    found = True
                    break
        if found:
            continue

        # Add remaining text as normal
        segments.append(('normal', text[current_pos:]))
        break

    # Add segments to the paragraph
    for style, content in segments:
        run = paragraph.add_run(content)
        if style == 'bold':
            run.bold = True
        elif style == 'italic':
            run.italic = True

def _format_description_word(description: str, document, level: int) -> None:
    """Format description text for Word, handling Markdown elements."""
    if not description:
        return

    for line in description.split('\n'):
        stripped_line = line.strip()
        if not stripped_line:
            continue

        # Check for headers
        if stripped_line.startswith('#'):
            md_header_level = 0
            while md_header_level < len(stripped_line) and stripped_line[md_header_level] == '#':
                md_header_level += 1
            if md_header_level < len(stripped_line) and stripped_line[md_header_level] == ' ':
                header_text = stripped_line[md_header_level+1:].strip()
                effective_level = level + md_header_level
                if effective_level >= 5:
                    style = 'Heading 5'
                else:
                    style = ['Title', 'Heading 1', 'Heading 2', 'Heading 3', 'Heading 4'][effective_level]
                document.add_paragraph(header_text, style=style)
                continue

        # Check for unordered list
        unordered_patterns = ['* ', '- ', '+ ']
        if any(stripped_line.startswith(p) for p in unordered_patterns):
            list_text = stripped_line[2:].strip()
            p = document.add_paragraph(style='List Bullet')
            _add_inline_formatting(p, list_text)
            continue

        # Check for ordered list
        ordered_match = re.match(r'^(\d+)\.\s+', stripped_line)
        if ordered_match:
            list_text = stripped_line[ordered_match.end():].strip()
            p = document.add_paragraph(style='List Number')
            _add_inline_formatting(p, list_text)
            continue

        # Regular paragraph with inline formatting
        p = document.add_paragraph()
        _add_inline_formatting(p, stripped_line)

def _process_node_word(document: Document, node: LayoutModel, level: int = 0) -> None:
    """Process a node and its children recursively to generate Word content."""
    # Add header with proper heading level
    if level == 0:
        if node.name:
            document.add_paragraph(node.name, style='Title')
    elif level == 1:
        if node.name:
            document.add_paragraph(node.name, style='Heading 1')
    elif level == 2:
        if node.name:
            document.add_paragraph(node.name, style='Heading 2')
    elif level == 3:
        if node.name:
            document.add_paragraph(node.name, style='Heading 3')
    elif level == 4:
        if node.name:
            document.add_paragraph(node.name, style='Heading 4')
    else:
        if node.name:
            document.add_paragraph(node.name, style='Heading 5')

    # Add description if present
    if node.description:
        _format_description_word(node.description, document, level)

    # Process children if present
    if node.children:
        for child in node.children:
            _process_node_word(document, child, level + 1)

def export_to_word(model: LayoutModel, settings: Settings) -> Document:
    """Export the capability model to a Microsoft Word Document object."""
    document = Document()

    # Add title page
    title_section = document.sections[0]
    title_paragraph = document.add_paragraph()
    title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = title_paragraph.add_run(model.name if model.name else "Document Title")
    run.bold = True
    run.font.size = Pt(24)

    subtitle_paragraph = document.add_paragraph()
    subtitle_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle_paragraph.add_run("Generated by Themis")
    run.font.size = Pt(16)

    date_paragraph = document.add_paragraph()
    date_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    today = datetime.date.today()
    run = date_paragraph.add_run(today.strftime("%d %B %Y"))
    run.font.size = Pt(12)

    document.add_page_break()

    processed_model = process_layout(model, settings)
    _process_node_word(document, processed_model, level=0)
    return document
